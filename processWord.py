from docx import Document
import csv
import sys

def sameStyle(paragraph1,paragraph2):
    if  (paragraph1['stylename'] == paragraph2['stylename'] and
        paragraph1['font_size'] == paragraph2['font_size'] and
        paragraph1['font_color'] == paragraph2['font_color']):
        return True
    else:
        return False
    
# 打开Word文件
doc = Document('data/xai.docx')

# 读取段落
content = []
for paragraph in doc.paragraphs:
    if paragraph.text !='':
        for run in paragraph.runs:
            content.append({
                "text": run.text,
                "font_size": run.font.size,
                "font_color": run.font.color.rgb if run.font.color else None,
                "stylename": paragraph.style.name
            })
if not content:
    print("No content found.")
    sys.exit(1)
    
# 合并相同样式的文本
new_content = [content[0]]
for i in range(1,len(content)):
    if sameStyle(content[i],new_content[-1]):
        new_content[-1]['text'] += content[i]['text']
    else:
        new_content.append(content[i]) 
         
# 读取表格
tables = []
for table in doc.tables:
        for row in table.rows:
            tstr='\t'.join(cell.text for cell in row.cells)
            tables.append({"table":tstr})
                 
# 保存为csv文件
with open('data\output.csv', 'w', newline='', encoding='utf-8-sig') as file:
    writer = csv.DictWriter(file, fieldnames=["text", "stylename","font_size","font_color","table"])
    writer.writeheader()
    writer.writerows(new_content)
    writer.writerows(tables)

    

